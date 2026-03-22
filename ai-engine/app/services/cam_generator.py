"""
CAM DOCX generator: produces a professional Credit Appraisal Memo as a .docx file.

Sections:
1. Executive Summary
2. Borrower Profile
3. Business Overview
4. Financial Analysis (table)
5. Risk Analysis
6. SWOT Analysis
7. Five Cs of Credit
8. Red Flags
9. Final Recommendation
10. Evidence Appendix
"""
from __future__ import annotations

import logging
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

DATA_ROOT = Path(os.environ.get("DATA_ROOT", "../data"))

# Try importing python-docx; graceful fallback to markdown if not available
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed; CAM will be generated as markdown fallback.")


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def _fmt_inr(value: float | None) -> str:
    if value is None:
        return "N/A"
    if abs(value) >= 1e7:
        return f"₹{value / 1e7:,.2f} Cr"
    if abs(value) >= 1e5:
        return f"₹{value / 1e5:,.2f} L"
    return f"₹{value:,.0f}"


def _get_v(facts: dict, key: str) -> Any:
    v = facts.get(key)
    if isinstance(v, dict):
        return v.get("value")
    return v


def _fmt_ratio(value: Any) -> str:
    """FIX #A1: 'value or N/A' wrongly converts 0.0 to N/A.
    Use explicit None check so a genuine 0.0 ratio is displayed."""
    if value is None:
        return "N/A"
    try:
        return f"{float(value):.2f}"
    except (TypeError, ValueError):
        return str(value)


def _severity_badge(sev: str) -> str:
    return {
        "critical": "🔴 CRITICAL",
        "high": "🟠 HIGH",
        "medium": "🟡 MEDIUM",
        "low": "🟢 LOW",
    }.get(sev.lower() if sev else "low", (sev or "").upper())


# FIX #A2: Explicit decision map that includes "reject" (previously missing).
DECISION_LABEL_MAP = {
    "approve": "APPROVE",
    "approve_with_conditions": "APPROVE WITH CONDITIONS",
    "decline": "DECLINE",
    "reject": "DECLINE",           # scoring_engine returns "reject"; map to label
    "manual_review": "MANUAL REVIEW",
}

# FIX #A3: Safe table style — "Table Grid" is always available in python-docx.
# "Light Grid Accent 1" requires the full Office theme to be loaded; if the
# document was created without it the style lookup raises KeyError and the
# whole DOCX generation crashes.
_TABLE_STYLE = "Table Grid"


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def _add_styled_heading(doc: "Document", text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)  # Dark navy


def _add_kv_line(doc: "Document", label: str, value: str):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_val = p.add_run(str(value))
    run_val.font.size = Pt(10)


def _add_financial_table(doc: "Document", facts: dict):
    rows_data = [
        ("Revenue",              _fmt_inr(_get_v(facts, "revenue"))),
        ("EBITDA",               _fmt_inr(_get_v(facts, "EBITDA"))),
        ("PAT (Profit After Tax)",_fmt_inr(_get_v(facts, "PAT"))),
        ("Total Debt",           _fmt_inr(_get_v(facts, "total_debt"))),
        ("Working Capital",      _fmt_inr(_get_v(facts, "working_capital"))),
        ("Current Ratio",        _fmt_ratio(_get_v(facts, "current_ratio"))),  # FIX #A1
        ("DSCR",                 _fmt_ratio(_get_v(facts, "dscr"))),           # FIX #A1
    ]
    table = doc.add_table(rows=len(rows_data) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = _TABLE_STYLE  # FIX #A3

    hdr = table.rows[0]
    hdr.cells[0].text = "Metric"
    hdr.cells[1].text = "Value"
    for cell in hdr.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    for i, (metric, val) in enumerate(rows_data, 1):
        table.rows[i].cells[0].text = metric
        table.rows[i].cells[1].text = str(val)


# ---------------------------------------------------------------------------
# SWOT & Five Cs computation
# ---------------------------------------------------------------------------

def _compute_swot(
    facts: dict,
    risk_flags: list[dict],
    score: dict,
    company: dict,
) -> dict[str, list[str]]:
    breakdown = score.get("score_breakdown", {})
    revenue = _get_v(facts, "revenue")
    dscr = _get_v(facts, "dscr")
    pat = _get_v(facts, "PAT")
    cr = _get_v(facts, "current_ratio")

    swot: dict[str, list[str]] = {
        "Strengths": [],
        "Weaknesses": [],
        "Opportunities": [],
        "Threats": [],
    }

    # --- Strengths ---
    if revenue and revenue > 1e7:
        swot["Strengths"].append(f"Established revenue base ({_fmt_inr(revenue)})")
    if pat is not None and pat > 0:
        swot["Strengths"].append("Profitable operations (positive PAT)")
    if cr is not None and cr >= 1.5:
        swot["Strengths"].append("Robust short-term liquidity (CR ≥ 1.5)")
    if dscr is not None and dscr >= 1.3:
        swot["Strengths"].append(f"Comfortable debt service coverage (DSCR {dscr:.2f})")
    if breakdown.get("governance", 0) > 85:
        swot["Strengths"].append("Strong governance indicators")
    if not swot["Strengths"]:
        swot["Strengths"].append("No significant financial strengths identified from available data.")

    # --- Weaknesses ---
    # FIX #A4: Collect weakness flag_types to prevent duplicating them in Threats.
    weakness_flag_types = {
        "negative_working_capital", "low_factory_utilization", "weak_cash_conversion",
    }
    if pat is not None and pat < 0:
        swot["Weaknesses"].append("Currently operating at a net loss")
    if cr is not None and cr < 1.0:
        swot["Weaknesses"].append(f"Strained short-term liquidity (CR {cr:.2f} < 1.0)")
    if dscr is not None and dscr < 1.0:
        swot["Weaknesses"].append(f"Inadequate debt service coverage (DSCR {dscr:.2f})")
    for f in risk_flags:
        if f.get("flag_type") in weakness_flag_types:
            desc = f.get("description", "")
            if desc and desc not in swot["Weaknesses"]:
                swot["Weaknesses"].append(desc)
    if not swot["Weaknesses"]:
        swot["Weaknesses"].append("No major financial weaknesses identified.")

    # --- Opportunities ---
    if company.get("sector"):
        swot["Opportunities"].append(
            f"Growth potential in the {company['sector']} sector"
        )
    # FIX #A5: Only add the "leverage track record" opportunity for non-risky borrowers.
    overall_score = score.get("overall_score", 0)
    decision = score.get("decision", "review")
    if overall_score >= 65 and decision not in ("reject", "decline"):
        swot["Opportunities"].append(
            "Opportunity to leverage positive track record for competitive financing terms"
        )
    else:
        swot["Opportunities"].append(
            "Potential to improve financial profile with operational restructuring"
        )
    if not swot["Opportunities"]:
        swot["Opportunities"].append("Sector and market opportunities to be assessed by analyst.")

    # --- Threats ---
    # FIX #A4: Track added descriptions to deduplicate.
    seen_threats: set[str] = set()

    def _add_threat(desc: str):
        if desc and desc not in seen_threats and desc not in swot["Weaknesses"]:
            swot["Threats"].append(desc)
            seen_threats.add(desc)

    for f in risk_flags:
        sev = f.get("severity", "low")
        ft = f.get("flag_type", "")
        if sev in ("high", "critical") and ft not in weakness_flag_types:
            _add_threat(f.get("description", ""))

    if breakdown.get("secondary_research", 100) < 70:
        _add_threat("Adverse signals from secondary web research or news")

    if not swot["Threats"]:
        swot["Threats"].append("No severe external threats or critical red flags identified.")

    return swot


def _compute_five_cs(
    facts: dict,
    flags: list[dict],
    score: dict,
    company: dict,
) -> dict[str, str]:
    breakdown = score.get("score_breakdown", {})
    revenue = _get_v(facts, "revenue")
    dscr = _get_v(facts, "dscr")
    cr = _get_v(facts, "current_ratio")

    gov_concern = any(
        f.get("flag_type") in ("governance_instability", "high_related_party_dependency")
        for f in flags
    )
    sector_concern = any(
        f.get("flag_type") in ("regulatory_risk", "sector_headwind")
        for f in flags
    )

    return {
        "Character": (
            f"Promoter(s): {', '.join(company.get('promoter_names', [])) or 'N/A'}. "
            f"Governance score: {breakdown.get('governance', 'N/A')}/100. "
            + ("Governance concerns flagged." if gov_concern else "No major governance concerns identified.")
        ),
        "Capacity": (
            f"DSCR: {_fmt_ratio(dscr)}. Cash flow score: {breakdown.get('cash_flow', 'N/A')}/100. "
            + (
                f"DSCR of {dscr:.2f} is {'adequate' if dscr >= 1.0 else 'insufficient'} for debt servicing."
                if dscr is not None
                else "DSCR not available — assess repayment capacity manually."
            )
        ),
        "Capital": (
            f"Revenue: {_fmt_inr(revenue)}. Financial strength score: {breakdown.get('financial_strength', 'N/A')}/100. "
            f"Working capital: {_fmt_inr(_get_v(facts, 'working_capital'))}."
        ),
        "Collateral": (
            "Collateral assessment based on uploaded documentation. "
            "Tangible security details should be verified from sanction letters and valuation reports."
        ),
        "Conditions": (
            f"Sector: {company.get('sector', 'N/A')}. Current ratio: {_fmt_ratio(cr)}. "
            f"Secondary research score: {breakdown.get('secondary_research', 'N/A')}/100. "
            + ("Regulatory or sector headwinds noted." if sector_concern else "No sector-level concerns flagged.")
        ),
    }


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------

def generate_cam_docx(
    case_id: str,
    company_details: dict | None,
    extracted_facts: dict,
    risk_flags: list[dict],
    score_result: dict,
    officer_notes: str | None = None,
) -> str:
    """Generate a DOCX (or markdown fallback) CAM file. Returns the file path."""
    evidence_dir = DATA_ROOT / "evidence" / case_id
    evidence_dir.mkdir(parents=True, exist_ok=True)

    if not HAS_DOCX:
        return _generate_cam_markdown(
            case_id, company_details, extracted_facts,
            risk_flags, score_result, officer_notes, evidence_dir,
        )

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    cd = company_details or {}
    company     = cd.get("company_name", "Unknown Company")
    sector      = cd.get("sector", "N/A")
    promoters   = ", ".join(cd.get("promoter_names", [])) or "N/A"
    cin         = cd.get("cin_optional", "N/A")

    decision       = score_result.get("cam_decision") or score_result.get("decision") or "manual_review"
    overall_score  = score_result.get("overall_score", 0)
    rec_limit      = score_result.get("recommended_limit", 0)
    rec_roi        = score_result.get("recommended_roi", 0)
    reasons        = score_result.get("reasons", [])
    breakdown      = score_result.get("score_breakdown", {})

    # ── Title ──
    title = doc.add_heading("Credit Appraisal Memorandum", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    subtitle = doc.add_paragraph(
        f"Case ID: {case_id} | Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # ── 1. Executive Summary ──
    _add_styled_heading(doc, "1. Executive Summary", 1)
    _add_kv_line(doc, "Decision", DECISION_LABEL_MAP.get(decision, decision.upper()))
    _add_kv_line(doc, "Overall Score", f"{overall_score}/100")
    _add_kv_line(doc, "Recommended Limit", _fmt_inr(rec_limit))
    _add_kv_line(doc, "Recommended ROI", f"{rec_roi}%")

    if score_result.get("hard_override_applied"):
        p = doc.add_paragraph()
        run = p.add_run(
            f"⚠️ HARD OVERRIDE: {score_result.get('hard_override_reason', 'Red flag triggered')}"
        )
        run.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    explanation = score_result.get("decision_explanation", "")
    if explanation:
        doc.add_paragraph(explanation)

    # ── 2. Borrower Profile ──
    _add_styled_heading(doc, "2. Borrower Profile", 1)
    _add_kv_line(doc, "Company Name", company)
    _add_kv_line(doc, "CIN", cin)
    _add_kv_line(doc, "Sector", sector)
    _add_kv_line(doc, "Promoter(s)", promoters)

    # ── 3. Business Overview ──
    _add_styled_heading(doc, "3. Business Overview", 1)
    doc.add_paragraph(
        f"{company} operates in the {sector} sector. "
        "This appraisal is based on uploaded financial statements, regulatory filings, "
        "and secondary research. The company's financial profile and risk indicators "
        "are summarized below."
    )

    # ── 4. Financial Analysis ──
    _add_styled_heading(doc, "4. Financial Analysis", 1)
    _add_financial_table(doc, extracted_facts)
    doc.add_paragraph("")

    if breakdown:
        doc.add_heading("Score Breakdown", level=2)
        tbl = doc.add_table(rows=len(breakdown) + 1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = _TABLE_STYLE  # FIX #A3
        tbl.rows[0].cells[0].text = "Category"
        tbl.rows[0].cells[1].text = "Score"
        for cell in tbl.rows[0].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
        for i, (cat, val) in enumerate(breakdown.items(), 1):
            tbl.rows[i].cells[0].text = cat.replace("_", " ").title()
            tbl.rows[i].cells[1].text = str(val)

    # ── 5. Risk Analysis ──
    _add_styled_heading(doc, "5. Risk Analysis", 1)
    if risk_flags:
        for flag in risk_flags:
            sev = flag.get("severity", "medium")
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{_severity_badge(sev)}] ")
            run.bold = True
            p.add_run(f"{flag.get('flag_type', 'unknown')}: {flag.get('description', '')}")
            refs = flag.get("evidence_refs", [])
            if refs:
                p.add_run(f" (Refs: {', '.join(str(r) for r in refs[:3])})")
    else:
        doc.add_paragraph("No significant risk flags identified.")

    # ── 6. SWOT Analysis ──
    _add_styled_heading(doc, "6. SWOT Analysis", 1)
    swot = _compute_swot(extracted_facts, risk_flags, score_result, cd)
    for category, items in swot.items():
        doc.add_heading(category, level=2)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    # ── 7. Five Cs of Credit ──
    _add_styled_heading(doc, "7. Five Cs of Credit", 1)
    five_cs = _compute_five_cs(extracted_facts, risk_flags, score_result, cd)
    for c_name, c_assessment in five_cs.items():
        doc.add_heading(c_name, level=2)
        doc.add_paragraph(c_assessment)

    # ── 8. Red Flags ──
    _add_styled_heading(doc, "8. Red Flags", 1)
    severe = [f for f in risk_flags if f.get("severity") in ("high", "critical")]
    if severe:
        for flag in severe:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{(flag.get('severity') or '').upper()}] ")
            run.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            p.add_run(flag.get("description", ""))
    else:
        doc.add_paragraph("No severe red flags identified.")

    if score_result.get("hard_override_applied"):
        p = doc.add_paragraph()
        run = p.add_run(f"Hard Override Active: {score_result.get('hard_override_reason', '')}")
        run.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # ── 9. Final Recommendation ──
    _add_styled_heading(doc, "9. Final Recommendation", 1)
    _add_kv_line(doc, "Decision", DECISION_LABEL_MAP.get(decision, decision.upper()))
    _add_kv_line(doc, "Overall Score", f"{overall_score}/100")
    _add_kv_line(doc, "Recommended Exposure Limit", _fmt_inr(rec_limit))
    _add_kv_line(doc, "Recommended ROI", f"{rec_roi}%")

    if reasons:
        doc.add_heading("Key Reasons", level=2)
        for r in reasons[:8]:
            doc.add_paragraph(r, style="List Bullet")

    if officer_notes:
        doc.add_heading("Officer Notes", level=2)
        doc.add_paragraph(officer_notes)

    # ── 10. Evidence Appendix ──
    _add_styled_heading(doc, "10. Evidence Appendix", 1)
    doc.add_paragraph("The following evidence sources were referenced in this appraisal:")

    entities = extracted_facts.get("extracted_entities", [])
    if isinstance(entities, list):
        for ent in entities[:10]:
            if isinstance(ent, dict):
                doc.add_paragraph(
                    f"• {ent.get('name', 'N/A')} ({ent.get('type', 'N/A')}) "
                    f"— {ent.get('role', '')} [Source: {ent.get('source_ref', 'N/A')}]"
                )

    doc_sources = extracted_facts.get("document_sources", [])
    if isinstance(doc_sources, list):
        for ds in doc_sources[:10]:
            if isinstance(ds, dict):
                doc.add_paragraph(
                    f"• {ds.get('file_name', 'N/A')} "
                    f"— Section: {ds.get('section', 'N/A')}, Page: {ds.get('page_ref', 'N/A')}"
                )

    # Save
    docx_path = evidence_dir / "cam.docx"
    doc.save(str(docx_path))
    logger.info("CAM DOCX generated: %s", docx_path)
    return str(docx_path)


# ---------------------------------------------------------------------------
# Markdown fallback when python-docx is not installed
# ---------------------------------------------------------------------------

def _generate_cam_markdown(
    case_id: str,
    company_details: dict | None,
    extracted_facts: dict,
    risk_flags: list[dict],
    score_result: dict,
    officer_notes: str | None,
    evidence_dir: Path,
) -> str:
    cd = company_details or {}
    company       = cd.get("company_name", "Unknown Company")
    decision      = score_result.get("cam_decision") or score_result.get("decision") or "manual_review"
    overall_score = score_result.get("overall_score", 0)
    rec_limit     = score_result.get("recommended_limit", 0)
    rec_roi       = score_result.get("recommended_roi", 0)
    reasons       = score_result.get("reasons", [])
    breakdown     = score_result.get("score_breakdown", {})

    lines = [
        "# Credit Appraisal Memorandum",
        f"**Case ID:** {case_id} | **Generated:** {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
        "",
        "## 1. Executive Summary",
        f"- **Decision:** {DECISION_LABEL_MAP.get(decision, decision.upper())}",
        f"- **Overall Score:** {overall_score}/100",
        f"- **Recommended Limit:** {_fmt_inr(rec_limit)}",
        f"- **Recommended ROI:** {rec_roi}%",
        "",
        "## 2. Borrower Profile",
        f"- **Company:** {company}",
        f"- **CIN:** {cd.get('cin_optional', 'N/A')}",
        f"- **Sector:** {cd.get('sector', 'N/A')}",
        f"- **Promoters:** {', '.join(cd.get('promoter_names', [])) or 'N/A'}",
        "",
        "## 3. Financial Analysis",
        "| Metric | Value |",
        "|--------|-------|",
        f"| Revenue | {_fmt_inr(_get_v(extracted_facts, 'revenue'))} |",
        f"| EBITDA | {_fmt_inr(_get_v(extracted_facts, 'EBITDA'))} |",
        f"| PAT | {_fmt_inr(_get_v(extracted_facts, 'PAT'))} |",
        f"| Total Debt | {_fmt_inr(_get_v(extracted_facts, 'total_debt'))} |",
        # FIX #A6: working_capital was in DOCX table but missing from markdown table
        f"| Working Capital | {_fmt_inr(_get_v(extracted_facts, 'working_capital'))} |",
        f"| Current Ratio | {_fmt_ratio(_get_v(extracted_facts, 'current_ratio'))} |",   # FIX #A1
        f"| DSCR | {_fmt_ratio(_get_v(extracted_facts, 'dscr'))} |",                     # FIX #A1
        "",
        "## 4. Risk Flags",
    ]

    for flag in risk_flags:
        lines.append(
            f"- **[{(flag.get('severity') or 'N/A').upper()}]** "
            f"{flag.get('flag_type', '')}: {flag.get('description', '')}"
        )
    if not risk_flags:
        lines.append("- No significant risk flags.")

    lines.extend(["", "## 5. SWOT Analysis"])
    swot = _compute_swot(extracted_facts, risk_flags, score_result, cd)
    for category, items in swot.items():
        lines.append(f"### {category}")
        for item in items:
            lines.append(f"- {item}")

    lines.extend([
        "",
        "## 6. Score Breakdown",
        "| Category | Score |",
        "|----------|-------|",
    ])
    for cat, val in breakdown.items():
        lines.append(f"| {cat.replace('_', ' ').title()} | {val} |")

    lines.extend([
        "",
        "## 7. Five Cs of Credit",
    ])
    five_cs = _compute_five_cs(extracted_facts, risk_flags, score_result, cd)
    for c_name, c_text in five_cs.items():
        lines.append(f"**{c_name}:** {c_text}")
        lines.append("")

    lines.extend([
        "## 8. Recommendation",
        f"**Decision:** {DECISION_LABEL_MAP.get(decision, decision.upper())} | **Score:** {overall_score}/100",
    ])
    if reasons:
        lines.append("### Key Reasons")
        for r in reasons[:8]:
            lines.append(f"- {r}")

    if officer_notes:
        lines.extend(["", "## 9. Officer Notes", officer_notes])

    lines.extend([
        "",
        "## 10. Evidence Appendix",
        "The following evidence sources were referenced in this appraisal:",
    ])
    entities = extracted_facts.get("extracted_entities", [])
    if isinstance(entities, list):
        for ent in entities[:10]:
            if isinstance(ent, dict):
                lines.append(
                    f"- {ent.get('name', 'N/A')} ({ent.get('type', 'N/A')}) "
                    f"— {ent.get('role', '')} [Source: {ent.get('source_ref', 'N/A')}]"
                )
    doc_sources = extracted_facts.get("document_sources", [])
    if isinstance(doc_sources, list):
        for ds in doc_sources[:10]:
            if isinstance(ds, dict):
                lines.append(
                    f"- {ds.get('file_name', 'N/A')} "
                    f"— Section: {ds.get('section', 'N/A')}, Page: {ds.get('page_ref', 'N/A')}"
                )

    md_path = evidence_dir / "cam.md"
    md_path.write_text("\n".join(lines), encoding="utf-8")
    logger.info("CAM markdown generated (python-docx fallback): %s", md_path)
    return str(md_path)